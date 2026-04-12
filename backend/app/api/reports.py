"""Report generation and download routes (v2 — PDF + DOCX support)."""

from __future__ import annotations

import logging
from datetime import datetime, timezone
from pathlib import Path
from uuid import uuid4

from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import FileResponse

from app.agents.report_writer import ReportWriterAgent
from app.config import settings
from app.models.schemas import ReportListItem, ReportRequest, ReportResponse
from app.models.user import User
from app.services.auth_service import get_current_user

logger = logging.getLogger(__name__)
router = APIRouter()


def _markdown_to_pdf(markdown_text: str, pdf_path: Path) -> None:
    """Render markdown-like plain text into a formatted PDF document."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas

    pdf = canvas.Canvas(str(pdf_path), pagesize=A4)
    text_object = pdf.beginText(40, 800)
    text_object.setFont("Helvetica", 11)

    for line in markdown_text.splitlines():
        if text_object.getY() < 40:
            pdf.drawText(text_object)
            pdf.showPage()
            text_object = pdf.beginText(40, 800)
            text_object.setFont("Helvetica", 11)
        clean = line.replace("**", "").replace("*", "").replace("#", "").strip()
        text_object.textLine(clean[:120])

    pdf.drawText(text_object)
    pdf.save()


def _markdown_to_docx(markdown_text: str, docx_path: Path) -> None:
    """Render markdown-like plain text into a Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Calibri"

        for line in markdown_text.splitlines():
            clean = line.strip()
            if not clean:
                doc.add_paragraph("")
                continue
            if clean.startswith("# "):
                doc.add_heading(clean[2:], level=1)
            elif clean.startswith("## "):
                doc.add_heading(clean[3:], level=2)
            elif clean.startswith("- "):
                doc.add_paragraph(clean[2:].replace("**", ""), style="List Bullet")
            else:
                doc.add_paragraph(clean.replace("**", "").replace("*", ""))

        doc.save(str(docx_path))
    except ImportError:
        logger.warning("python-docx not installed, falling back to plain text")
        docx_path.with_suffix(".txt").write_text(markdown_text)


@router.post("/generate", response_model=ReportResponse)
async def generate_report(
    payload: Optional[ReportRequest] = None,
    current_user: User = Depends(get_current_user),
) -> ReportResponse:
    """Generate and persist a compliance report in the requested format."""
    request = payload or ReportRequest()
    writer = ReportWriterAgent()
    result = await writer.generate_compliance_report(str(current_user.id))
    content = result["content"]
    report_id = uuid4().hex
    report_path = settings.reports_dir / f"{report_id}.md"
    report_path.write_text(content, encoding="utf-8")

    if request.format == "pdf":
        _markdown_to_pdf(content, settings.reports_dir / f"{report_id}.pdf")
    elif request.format == "docx":
        _markdown_to_docx(content, settings.reports_dir / f"{report_id}.docx")

    return ReportResponse(
        report_id=report_id,
        content=content,
        report_type=request.report_type,
        generated_at=datetime.now(timezone.utc),
        citations=result.get("citations", []),
        confidence_score=result.get("confidence", 1.0),
    )


@router.get("/list", response_model=list[ReportListItem])
async def list_reports(current_user: User = Depends(get_current_user)) -> list[ReportListItem]:
    """List all generated reports."""
    _ = current_user
    reports: list[ReportListItem] = []
    if settings.reports_dir.exists():
        for md_file in sorted(settings.reports_dir.glob("*.md"), key=lambda p: p.stat().st_mtime, reverse=True):
            report_id = md_file.stem
            stat = md_file.stat()
            reports.append(ReportListItem(
                report_id=report_id,
                report_type="compliance",
                generated_at=datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc),
                filename=f"cubitaxai-report-{report_id}",
            ))
    return reports[:20]


@router.get("/download/{report_id}")
async def download_report(
    report_id: str,
    format: str = "pdf",
    current_user: User = Depends(get_current_user),
) -> FileResponse:
    """Return a generated report as a downloadable PDF or DOCX."""
    _ = current_user
    markdown_path = settings.reports_dir / f"{report_id}.md"
    if not markdown_path.exists():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")

    if format == "docx":
        docx_path = settings.reports_dir / f"{report_id}.docx"
        if not docx_path.exists():
            _markdown_to_docx(markdown_path.read_text(encoding="utf-8"), docx_path)
        return FileResponse(
            path=docx_path,
            filename=f"cubitaxai-report-{report_id}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    pdf_path = settings.reports_dir / f"{report_id}.pdf"
    if not pdf_path.exists():
        _markdown_to_pdf(markdown_path.read_text(encoding="utf-8"), pdf_path)
    return FileResponse(path=pdf_path, filename=f"cubitaxai-report-{report_id}.pdf", media_type="application/pdf")
